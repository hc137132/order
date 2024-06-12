import json
import time
from django.test import TestCase
import os
import datetime
import uuid
import re
import io
from docx.shared import Pt, Inches
from tool.pill import getsql

# Create your tests here.
current_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
path=os.path.join(current_dir,r'static\files',str(datetime.date.today()))
print(str(uuid.uuid1()).replace('-',''))
# if os.path.exists(path + '/'+str(datetime.date.today())):
#     print('yes',path)
# else:
#     os.mkdir(path + '/'+str(datetime.date.today()))
#     print('no',)
def a(a,b):
    if a and b:
      pass
    else:
        a=1
        b=2
    print(a, b)

# a()
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
def word(docx_path):
    # 打开word文档
    doc = Document(docx_path)
    # 获取所有段落
    # all_paragraphs = document.part._rels
    # 创建一个新的Word文档，用于存储提取的内容
    new_doc = Document(r'D:\桌面\Work\医院\2021.docx')
    dict_rel = doc.part._rels
    for paragraph in doc.paragraphs:
        new_doc.add_paragraph(paragraph.text)
        # run = paragraph.runs[0]
        # run.font.size = Pt(12)  # 设置字体大小为12磅
        #
        # # 设置段落的对齐方式为居中
        # paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #     # 打印每一个段落的文字
    #     print(all_paragraphs[paragraph].target_ref)
    #     # 循环读取每个段落里的run内容
    #     # for run in paragraph.runs:
    #     #     print(run)  # 打印run内容

    dict_rel = doc.part._rels  # rels其实是个目录
    # 创建一个新的段落
    new_paragraph = new_doc.add_paragraph()

    # 添加图片到新段落中
    new_run = new_paragraph.add_run()
    for rel in dict_rel:
        rel = dict_rel[rel]
        if "image" in rel.target_ref:
            img_name = re.findall("/(.*)", rel.target_ref)[0]
            # print(img_name)
            new_run.add_picture(io.BytesIO(rel.target_part.blob),Inches(2.7))
            # if len(doc.paragraphs) > 0 and doc.paragraphs[-1].text == "":
            #     doc.paragraphs[-1].clear()

    new_doc.save(r'D:\桌面\Work\医院\2021.docx')

# word('D:\桌面\Work\医院\上报卫健局\阆中市人民医院安全大排查大整治安全培训及消防演练.docx')
# word('D:\桌面\你好.docx')
# print(os.walk(r'D:\桌面\Work\医院\2021'))
# for a,b,c in os.walk(r'D:\桌面\Work\医院\2021'):
#     for name in c:
#         word(r'D:\桌面\Work\医院\2021'+'\\'+name)
#         print('已添加',name)
b=17
a='12天'
print(re.findall(r'\d+',a))

def getsql1(filterdict):
    sql='select * from app_projectdetail where +  -  @ /;'
    if filterdict['type']:
        print('s')
        sql=sql.split('+')[0] + 'type in {}'.format(tuple(filterdict['type'])) + sql.split('+')[1]
        print(sql)
    else:
        sql='select * from app_projectdetail where type is not null - @ /;'
        print(sql)
    print(1712332800,time.time(),'sada'.count('-'),'20-10'.split('-'))
    onedaytime=86400#一天的时间戳，单位是秒

    # first 当前时间戳
    nowtime =int(time.time())

    header = sql.split('-')[0]
    end = sql.split('-')[1]
    if filterdict['deadline']:
        n=0
        deadlinesyn=''
        for a in filterdict['deadline']:
            if n==0:#first header not add and/or
                n+=1
                if a.count('-')==1:#first add sql where,前面+and
                    gt=nowtime+onedaytime*int(a.split('-')[0])#现在时刻+结束的天数的时间戳，如 2024-4-3的时间错+30天的时间错=2024-5-2的时间戳
                    lt=nowtime+onedaytime*int(a.split('-')[1])#现在时刻+结束的天数的时间戳，如 2024-4-3的时间错+20天的时间错=2024-4-23的时间戳
                    deadlinesyn=deadlinesyn+'deadline between {b} and {a}'.format(a=gt,b=lt)#start dont gt end
                elif a=='7天以下':
                    gt = nowtime + onedaytime * 7
                    deadlinesyn=deadlinesyn+'deadline between 200 and {a}'.format(a=gt)
                elif a =='商议':

                    deadlinesyn=deadlinesyn+'deadline <= 200'
                elif a=='30天以上':
                    gt = nowtime + onedaytime * 30
                    deadlinesyn = deadlinesyn + 'deadline >= {a}'.format(a=gt, )

            else:#second header add or
                if a.count('-')==1:#first add sql where,前面+and
                    gt=nowtime+onedaytime*int(a.split('-')[0])#现在时刻+结束的天数的时间戳，如 2024-4-3的时间错+30天的时间错=2024-5-2的时间戳
                    lt=nowtime+onedaytime*int(a.split('-')[1])#现在时刻+结束的天数的时间戳，如 2024-4-3的时间错+20天的时间错=2024-4-23的时间戳
                    deadlinesyn=deadlinesyn+' or deadline between {b} and {a}'.format(a=gt,b=lt)
                    print(sql,n)
                elif a=='7天以下':
                    gt = nowtime + onedaytime * 7
                    deadlinesyn=deadlinesyn+' or deadline between 200 and {a}'.format(a=gt)
                elif a =='商议':

                    deadlinesyn=deadlinesyn+' or deadline <= 200'
                elif a=='30天以上':
                    gt = nowtime + onedaytime * 30
                    deadlinesyn = deadlinesyn + ' or deadline >= {a}'.format(a=gt, )


        sql=header+'and ( '+deadlinesyn+' )'+end
        # print(sql.split('+')[0] + 'type in {}'.format(tuple(filterdict['type'])) + sql.split('+')[1])
    else:
        sql=header+'and deadline is not null'+end
        print(sql)
    header = sql.split('@')[0]
    end = sql.split('@')[1]
    if filterdict['createdate']:
        n=0
        createdatesyn=''
        for a in filterdict['createdate']:
            # print(a)
            if n==0:#first header not add and/or
                n+=1
                if a.count('-')==1:#first add sql where,前面+and
                    gt=nowtime-onedaytime*int(a.split('-')[0])#现在时刻+结束的天数的时间戳，如 2024-4-3的时间错+30天的时间错=2024-5-2的时间戳
                    lt=nowtime-onedaytime*int(re.findall(r'\d+',a.split('-')[1])[0])#现在时刻+结束的天数的时间戳，如 2024-4-3的时间错+20天的时间错=2024-4-23的时间戳,后面可能含有文字，要删掉文字
                    createdatesyn=createdatesyn+'createdate between {a} and {b}'.format(a=gt,b=lt)#start dont gt end
                elif a=='7天以下':
                    gt = nowtime - onedaytime * 7
                    createdatesyn=createdatesyn+'createdate >= {a}'.format(a=gt)

                elif a=='30天以上':
                    gt = nowtime - onedaytime * 30
                    createdatesyn = createdatesyn + 'createdate <= {a}'.format(a=gt, )

            else:#second header add or
                if a.count('-')==1:#first add sql where,前面+and
                    gt=nowtime-onedaytime*int(a.split('-')[0])#现在时刻+结束的天数的时间戳，如 2024-4-3的时间错+30天的时间错=2024-5-2的时间戳
                    lt=nowtime-onedaytime*int(re.findall(r'\d+',a.split('-')[1])[0])#现在时刻+结束的天数的时间戳，如 2024-4-3的时间错+20天的时间错=2024-4-23的时间戳
                    createdatesyn=createdatesyn+' or createdate between {a} and {b}'.format(a=gt,b=lt)

                elif a=='7天以下':
                    gt = nowtime - onedaytime * 7
                    createdatesyn=createdatesyn+' or createdate >={a}'.format(a=gt)
                elif a=='30天以上':
                    gt = nowtime - onedaytime * 30
                    createdatesyn = createdatesyn + ' or createdate <= {a}'.format(a=gt, )

            # print(createdatesyn)
        sql=header+'and ( '+createdatesyn+' )'+end
        # print(sql.split('+')[0] + 'type in {}'.format(tuple(filterdict['type'])) + sql.split('+')[1])
    else:
        sql=header+' and createdate is not null '+end
        print(sql)
    print(sql)

    header = sql.split('/')[0]
    end = sql.split('/')[1]
    if filterdict['money']:
        n = 0
        moneysyn = ''
        for a in filterdict['money']:
            # print(a)
            if n == 0:  # first header not add and/or
                n += 1
                if a.count('-') == 1:  # first add sql where,前面+and
                    gt = int(a.split('-')[0])  # 现在时刻+结束的天数的时间戳，如 2024-4-3的时间错+30天的时间错=2024-5-2的时间戳
                    lt = int(a.split('-')[1]) # 现在时刻+结束的天数的时间戳，如 2024-4-3的时间错+20天的时间错=2024-4-23的时间戳,后面可能含有文字，要删掉文字
                    moneysyn = moneysyn + 'money between {b} and {a}'.format(a=gt, b=lt)  # start dont gt end
                elif a == '商议':
                    gt = nowtime - onedaytime * 7
                    moneysyn = moneysyn + 'money = 0'
                elif a == '5000元以上':
                    moneysyn = moneysyn + 'money >= 5000'
                elif a == '1000元以下':
                    moneysyn = moneysyn + 'money <= 1000'

            else:  # second header add or
                if a.count('-') == 1:  # first add sql where,前面+and
                    gt = int(a.split('-')[0])  # 现在时刻+结束的天数的时间戳，如 2024-4-3的时间错+30天的时间错=2024-5-2的时间戳
                    lt = int(a.split('-')[1])  # 现在时刻+结束的天数的时间戳，如 2024-4-3的时间错+20天的时间错=2024-4-23的时间戳
                    moneysyn = moneysyn + ' or money between {b} and {a}'.format(a=gt, b=lt)

                elif a == '商议':
                    moneysyn = moneysyn + ' or money = 0'
                elif a == '5000元以上':
                    moneysyn = moneysyn + ' or money >= 5000'
                elif a == '1000元以下':
                    moneysyn = moneysyn + ' or money <= 1000'

            # print(createdatesyn)
        sql = header + 'and ( ' + moneysyn + ' )' + end
    else:
        sql=header+';'

    print(sql)
    # sql="select * from app_projectdetail where type in ('asdas', 'ereretr') and   deadline = None or deadline >= 1717773083 or deadline between 1717773083 and 1716909083 or deadline between 1716909083 and 1716045083 or deadline between 1716045083 and 1715613083 or deadline <= 1715785883  * /"

dict1 = {
    'type': ['sda'], 'deadline': ['商议','30天以上','30-20','20-10','10-7','7天以下'],
    'createdate': ['30天以上','30-15天','15-7天','7天以下'], 'money': [ '商议','5000元以上','3000-2000','2000-1000','1000元以下']
}
print(getsql(dict1))